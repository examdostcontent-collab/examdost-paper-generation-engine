#!/usr/bin/env python3
r"""
inject_codes_docx.py  -- Question Coder for the "Boxed" .docx format

The Boxed paper is a .docx with ONE 2-column table per question. Column 0 holds row
labels; column 1 holds values. The two cells this script touches:

    | Question | <number> |   <- old per-test number (1..N); updated to the NEW number
    ...
    | OTS ID   |          |   <- currently blank; the question CODE is written here

For one test it filters the coding sheet by the Test-Number column, builds an
old-number -> (new-number, code) map, then for every question table:
  1. sets the OTS ID value cell to the Code
  2. (unless --no-renumber) sets the Question value cell to the new number

DEFAULT IS A DRY RUN. Pass --apply to write a NEW <name>_coded.docx (original kept).

Examples
--------
  # dry-run audit:
  python inject_codes_docx.py --excel sheet.csv --test "PGCIL EE 02" \
         --docx "Mock 02 Boxed.docx"

  # apply -> writes "Mock 02 Boxed_coded.docx" + audit:
  python inject_codes_docx.py --excel sheet.csv --test "PGCIL EE 02" \
         --docx "Mock 02 Boxed.docx" --apply
"""
import argparse
from pathlib import Path

import coder_common as cc


def cell_text(cell):
    return cell.text.strip()


def set_cell_text(cell, value):
    """Replace a cell's text while preserving the first run's formatting.

    python-docx cells hold paragraphs->runs. We write `value` into the first run of
    the first paragraph and clear the rest, so the cell's style/borders survive.
    """
    value = "" if value is None else str(value)
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = value
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.add_run(value)
    # clear any extra paragraphs in the cell
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""


def find_label_row(table, label):
    """Return the row whose col-0 text equals `label` (case-insensitive), else None."""
    target = label.strip().lower()
    for row in table.rows:
        if len(row.cells) >= 2 and cell_text(row.cells[0]).lower() == target:
            return row
    return None


def main():
    ap = argparse.ArgumentParser(description="Inject codes into OTS ID + renumber Question cells in a Boxed .docx.")
    ap.add_argument("--excel", required=True, help="Coding sheet (.csv/.xlsx)")
    ap.add_argument("--test", default=None,
                    help='Filter rows by this test value, e.g. "PGCIL EE 02". '
                         'Omit when the whole sheet is one document (subject-wise compiled).')
    ap.add_argument("--sheet", default=None,
                    help="Worksheet name/index for a multi-sheet .xlsx (e.g. a subject sheet)")
    ap.add_argument("--docx", required=True, help="The Boxed .docx to update")
    ap.add_argument("--out", help="Output path (default: <stem>_coded.docx)")
    ap.add_argument("--question-label", default="Question", help='Row label for the number cell')
    ap.add_argument("--ots-label", default="OTS ID", help='Row label for the code cell')
    ap.add_argument("--no-renumber", action="store_true", help="Only fill OTS ID; keep numbers")
    ap.add_argument("--overwrite-ots", action="store_true",
                    help="Overwrite OTS ID cells that are already non-blank (default: skip + warn)")
    ap.add_argument("--old-col", default=cc.DEFAULT_OLD_COL)
    ap.add_argument("--new-col", default=cc.DEFAULT_NEW_COL)
    ap.add_argument("--code-col", default=cc.DEFAULT_CODE_COL)
    ap.add_argument("--test-col", default=cc.DEFAULT_TEST_COL)
    ap.add_argument("--apply", action="store_true", help="Write output (default: dry run)")
    args = ap.parse_args()

    import docx

    excel, docx_path = Path(args.excel), Path(args.docx)
    if not excel.exists():
        raise SystemExit(f"ERROR: excel not found: {excel}")
    if not docx_path.exists():
        raise SystemExit(f"ERROR: docx not found: {docx_path}")

    rows, columns = cc.load_rows(excel, sheet=args.sheet)
    old_col = cc.resolve_col(args.old_col, columns, role="old number")
    new_col = cc.resolve_col(args.new_col, columns, role="new number")
    code_col = cc.resolve_col(args.code_col, columns, role="code")
    test_col = None
    if args.test is not None:
        test_col = cc.resolve_col(args.test_col, columns, fallback_to_last=True, role="test number")

    mapping, problems, matched_rows = cc.build_map(
        rows, old_col, new_col, code_col, test_col, args.test
    )
    if matched_rows == 0:
        msg = f"ERROR: no usable rows in sheet {args.sheet!r}."
        if args.test is not None:
            msg = (f"ERROR: no rows where {test_col!r} == {args.test!r}.\n"
                   f"  Available test values: {cc.available_tests(rows, test_col)}")
        raise SystemExit(msg)

    d = docx.Document(str(docx_path))

    found_old = []
    notes = []
    no_qrow = no_otsrow = ots_nonblank = renamed = coded = 0

    for ti, table in enumerate(d.tables):
        qrow = find_label_row(table, args.question_label)
        otsrow = find_label_row(table, args.ots_label)
        if qrow is None:
            no_qrow += 1
            continue  # not a question table
        old = cc.as_int_str(cell_text(qrow.cells[1]))
        if old == "":
            continue
        found_old.append(old)
        info = mapping.get(old)
        if info is None:
            continue  # reported as "in doc not in excel"

        # OTS ID
        if otsrow is None:
            no_otsrow += 1
        else:
            cur = cell_text(otsrow.cells[1])
            if cur and not args.overwrite_ots:
                ots_nonblank += 1
            else:
                if args.apply:
                    set_cell_text(otsrow.cells[1], info["code"])
                coded += 1

        # Question number
        if not args.no_renumber and info["new"] and info["new"] != old:
            if args.apply:
                set_cell_text(qrow.cells[1], info["new"])
            renamed += 1

    notes.append(f"## Document scan")
    notes.append(f"- Question tables found: **{len(found_old)}**")
    notes.append(f"- OTS ID cells to fill (code): **{coded}**")
    notes.append(f"- Question cells to renumber: **{renamed}**")
    if no_otsrow:
        notes.append(f"- ⚠️ tables with a Question row but NO '{args.ots_label}' row: **{no_otsrow}**")
    if ots_nonblank:
        notes.append(f"- ⚠️ OTS ID already non-blank (SKIPPED — use --overwrite-ots): **{ots_nonblank}**")
    notes.append(f"- non-question tables skipped: {no_qrow}")

    out_path = Path(args.out) if args.out else docx_path.with_name(docx_path.stem + "_coded.docx")
    audit_path = out_path.with_suffix(".audit.md")

    if args.apply:
        d.save(str(out_path))

    label = args.test or args.sheet or docx_path.stem
    report = cc.write_audit(
        audit_path, "Question Coder (docx)", label, mapping, found_old,
        problems, matched_rows, args.apply, out_path, extra_notes=notes,
    )
    print(report)
    print(f"\n[audit written to] {audit_path}")
    if args.apply:
        print(f"[coded docx written to] {out_path}")
    else:
        print("\nDRY RUN — re-run with --apply to write the coded .docx.")


if __name__ == "__main__":
    main()
