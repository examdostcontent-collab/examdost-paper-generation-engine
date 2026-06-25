"""
extract_text.py — dump the readable text of a question set so the classifier can
read every question regardless of the source format.

Usage:
    python extract_text.py <file>            # prints extracted text to stdout
    python extract_text.py <file> -o out.txt # also writes it to a file

Supports .docx, .pdf, .xlsx/.xlsm, .csv/.tsv, .txt/.md (and falls back to a
plain UTF-8 read for anything else). The goal is faithful, lightly-structured
text — paragraphs and table rows for Word, page-by-page text for PDF, one line
per row for spreadsheets — NOT a re-layout. The model then reads this and pulls
out each question's statement, options, answer, marks, and any existing tags.

For a `paper.json` (paper-generator output) you do not need this — feed that
straight to build_metadata_xlsx.py.

Dependencies (only the one for your format is needed):
    .docx -> python-docx ; .pdf -> pdfplumber (preferred) or PyMuPDF ;
    .xlsx -> openpyxl. Plain text/CSV need nothing.
"""
from __future__ import annotations

import sys
from pathlib import Path


def from_docx(path):
    import docx  # python-docx
    doc = docx.Document(str(path))
    out = []
    # paragraphs in document order
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    # tables (many question banks are laid out as tables)
    for ti, tbl in enumerate(doc.tables, start=1):
        out.append(f"\n[Table {ti}]")
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                out.append(" | ".join(cells))
    return "\n".join(out)


def from_pdf(path):
    # pdfplumber first (best at columns/spacing), fall back to PyMuPDF
    try:
        import pdfplumber
        out = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                out.append(f"\n[Page {i}]")
                out.append(page.extract_text() or "")
        return "\n".join(out)
    except ImportError:
        pass
    import fitz  # PyMuPDF
    out = []
    doc = fitz.open(str(path))
    for i, page in enumerate(doc, start=1):
        out.append(f"\n[Page {i}]")
        out.append(page.get_text())
    return "\n".join(out)


def from_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    out = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        out.append(f"\n[Sheet: {sheet}]")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                out.append(" | ".join(cells))
    wb.close()
    return "\n".join(out)


def from_csv(path):
    import csv
    out = []
    with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
        sample = f.read(4096)
        f.seek(0)
        delim = "\t" if (sample.count("\t") > sample.count(",")) else ","
        for row in csv.reader(f, delimiter=delim):
            if any(str(c).strip() for c in row):
                out.append(" | ".join(row))
    return "\n".join(out)


def extract(path):
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".docx":
        return from_docx(p)
    if ext == ".pdf":
        return from_pdf(p)
    if ext in (".xlsx", ".xlsm"):
        return from_xlsx(p)
    if ext in (".csv", ".tsv"):
        return from_csv(p)
    # .txt, .md, or unknown — plain read
    return p.read_text(encoding="utf-8", errors="replace")


def main():
    args = list(sys.argv[1:])
    out_path = None
    if "-o" in args:
        i = args.index("-o")
        out_path = args[i + 1]
        del args[i:i + 2]
    if len(args) != 1:
        print(__doc__)
        sys.exit(1)
    text = extract(args[0])
    if out_path:
        Path(out_path).write_text(text, encoding="utf-8")
        print(f"Wrote {out_path}  ({len(text)} chars)")
    else:
        sys.stdout.buffer.write(text.encode("utf-8", errors="replace"))


if __name__ == "__main__":
    main()
