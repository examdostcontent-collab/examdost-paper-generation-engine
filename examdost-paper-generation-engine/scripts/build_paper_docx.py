"""
build_paper_docx.py — render a generated paper (paper.json) as a branded Word
question bank.

Usage:
    python build_paper_docx.py <paper.json> <output.docx>

Honours meta.layout ("combined" | "split"), meta.metadata_tags, meta.show_solutions,
and renders [GEMINI_FLASH_PROMPT: ...] tags as labelled placeholder boxes for the
downstream image tool. Schema: references/paper_spec_schema.md.

Dependencies: python-docx.
"""
from __future__ import annotations

import json
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

import paperlib as P
import omml as OM
import latexmode as LX

_OMML_CACHE = {}   # {latex: omml_xml} populated per build() run

# LaTeX output mode: when True, _emit writes equations as inline LaTeX ($\scriptstyle ..$)
# text instead of native Word OMML — so no pandoc and no OMML->LaTeX conversion are needed.
# Set by a builder before build() (e.g. build_paper_boxed_docx --latex). MATH_SIZE is the
# inline size wrapper (\scriptstyle default; "" for full size).
LATEX_MODE = False
MATH_SIZE = LX.DEFAULT_SIZE

# ----------------------------- palette ------------------------------------
TEAL = RGBColor(0x0F, 0x8B, 0x8D)
NAVY = RGBColor(0x1F, 0x2A, 0x44)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
GREY = RGBColor(0x6B, 0x72, 0x80)
INK = RGBColor(0x2A, 0x2F, 0x3A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FILL_TEAL = "0F8B8D"
FILL_NAVY = "1F2A44"
FILL_TINT_TEAL = "E6F4F4"
FILL_TINT_GREY = "F4F4F6"
FILL_DIAGRAM = "FFF4E0"   # amber tint for diagram placeholders
FILL_WILD = "FCEFD9"


# ----------------------------- low-level helpers --------------------------
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _run(p, text, *, bold=False, italic=False, color=INK, size=10.5, font="Calibri"):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r.font.size = Pt(size)
    r.font.name = font
    return r


def _para(doc, *, space_before=0, space_after=4, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.08
    if indent is not None:
        pf.left_indent = Inches(indent)
    return p


def _box(doc, label, text, fill, label_color=NAVY):
    """A single-cell shaded box (used for diagram placeholders / notes)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, fill)
    cell.width = Inches(6.4)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if label:
        _run(p, f"{label}  ", bold=True, color=label_color, size=9)
    _run(p, text, italic=True, color=INK, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _diagram_boxes(doc, prompts):
    for tag in prompts or []:
        _box(doc, "[ DIAGRAM — Gemini prompt ]", P.gemini_prompt(tag), FILL_DIAGRAM, label_color=AMBER)


def _embed_image(doc, path, width_in=4.3, caption="Figure"):
    """Embed an actual rendered diagram image (PNG) inline, centred, with a caption."""
    if not path or not os.path.exists(path):
        return
    try:
        doc.add_picture(path, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = _para(doc, space_before=1, space_after=3)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cap, caption, italic=True, color=GREY, size=8.5)
    except Exception:
        pass


# ----------------------------- header / sections --------------------------
def add_header(doc, meta):
    # brand kicker
    p = _para(doc, space_after=0)
    _run(p, "EXAMDOST  ·  GENERATED QUESTION BANK", bold=True, color=TEAL, size=9)
    # title
    p = _para(doc, space_after=2)
    _run(p, P.render_math(meta.get("paper_title") or meta.get("exam", "Question Paper")),
         bold=True, color=NAVY, size=20)
    # meta line
    bits = []
    if meta.get("exam"):
        bits.append(meta["exam"])
    if meta.get("total_questions"):
        bits.append(f"{meta['total_questions']} Questions")
    if meta.get("total_marks"):
        bits.append(f"{meta['total_marks']} Marks")
    if meta.get("duration_min"):
        bits.append(f"{meta['duration_min']} min")
    if meta.get("organising_institute"):
        bits.append(f"Organising: {meta['organising_institute']}")
    if meta.get("generated_on"):
        bits.append(f"Generated {meta['generated_on']}")
    p = _para(doc, space_after=6)
    _run(p, "    |    ".join(str(b) for b in bits), color=GREY, size=9.5)

    instr = meta.get("instructions") or []
    if instr:
        p = _para(doc, space_before=2, space_after=2)
        _run(p, "Instructions", bold=True, color=TEAL, size=11)
        for line in instr:
            ip = _para(doc, space_after=1, indent=0.2)
            _run(ip, "•  ", bold=True, color=TEAL, size=10)
            _run(ip, P.render_math(str(line)), color=INK, size=10)
    _rule(doc)


def _rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), FILL_TEAL)
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(6)


def add_section_band(doc, title):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade(cell, FILL_NAVY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"  {title}", bold=True, color=WHITE, size=13)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ----------------------------- math emission ------------------------------
# Literal "\n" used as an authored line break (e.g. statement / assertion-reason
# stems: "...:\nI. ...\nII. ..."). Only a backslash-n before an UPPERCASE letter,
# a digit, or whitespace is a break — \nabla, \nu, \neq, \nonumber (lowercase) are
# real commands and must be left untouched.
_LIT_BREAK = re.compile(r"\\n(?=[A-Z0-9]|\s)")


def _split_lines(text):
    """Turn authored literal '\\n' separators into real line breaks; return lines."""
    return _LIT_BREAK.sub("\n", str(text)).split("\n")


def _emit(p, text, color=INK, size=10.5, bold=False, whole_math=False):
    """Emit text into paragraph p, turning math markup into NATIVE Word equations
    (OMML) where available, else falling back to Unicode (render_math). Authored
    '\\n' separators become real line breaks. In LATEX_MODE, emit inline LaTeX."""
    text = "" if text is None else str(text)
    if not text:
        return
    for li, line in enumerate(_split_lines(text)):
        if li:
            p.add_run().add_break()          # authored line break
        if not line:
            continue
        chunks = [("math", line)] if (whole_math and OM.is_equation(line)) else OM.split_math(line)
        if LATEX_MODE:
            for kind, val in chunks:
                if kind == "math":
                    _run(p, LX.inline(val, MATH_SIZE), bold=bold, color=color, size=size)
                elif val:
                    _run(p, val, bold=bold, color=color, size=size)   # prose stays literal
            continue
        for kind, val in chunks:
            if kind == "math":
                els = OM.omath_elements(_OMML_CACHE.get(val, "")) if _OMML_CACHE.get(val) else []
                if els:
                    for el in els:
                        p._p.append(el)
                    continue
                _run(p, P.render_math(val), bold=bold, color=color, size=size)   # fallback
            else:
                rendered = P.render_math(val)
                if rendered:
                    _run(p, rendered, bold=bold, color=color, size=size)


def _collect_math(spec):
    """All math strings a build will need, so they convert in one pandoc batch.
    Mirrors exactly what _emit looks up: whole formula lines + inline spans."""
    out = []

    def add_text(t):
        if t:
            for line in _split_lines(t):                       # mirror _emit's line split
                out.extend(v for k, v in OM.split_math(line) if k == "math")

    def add_whole(t):
        if t and str(t).strip():
            for line in _split_lines(t):
                if not line:
                    continue
                if OM.is_equation(line):
                    out.append(line)    # full equation line
                else:
                    out.extend(v for k, v in OM.split_math(line) if k == "math")

    for sec in spec.get("sections", []):
        for q in sec.get("questions", []):
            add_text(q.get("text"))
            for o in q.get("options") or []:
                add_text(o)
            add_text(q.get("concept"))
            add_text(q.get("answer"))
            sol = q.get("solution") or {}
            add_text(sol.get("concept"))
            for g in sol.get("given") or []:
                add_text(g)
            for f in sol.get("formula") or []:
                add_whole(f)                       # formula line = full equation
            for w in sol.get("where") or []:
                add_whole(w.get("sym"))
            for c in sol.get("calculation") or []:
                add_whole(c)                       # calc line = full equation (stacked)
            for o in sol.get("option_analysis") or []:
                add_text(o.get("text"))
            add_text(sol.get("final_answer"))
    return out


# ----------------------------- question rendering -------------------------
def _meta_line(doc, q, tags):
    if not tags:
        return
    parts = []
    label_map = {"Subject": "subject", "Chapter": "chapter", "Topic": "topic",
                 "Marks": "marks", "Type": "type"}
    for tag in tags:
        key = label_map.get(tag)
        if key and q.get(key) not in (None, ""):
            parts.append(f"{tag}: {q.get(key)}")
    if not parts:
        return
    p = _para(doc, space_after=2, indent=0.3)
    _run(p, "   ".join(parts), italic=True, color=GREY, size=8.5)


def render_stem(doc, q):
    """Question number + marks badge + text + options + (question) diagrams + metadata is
    handled by caller. This renders the stem and options only."""
    p = _para(doc, space_before=4, space_after=2)
    _run(p, f"Q{q.get('number', '')}.  ", bold=True, color=TEAL, size=11)
    if q.get("is_wildcard"):
        _run(p, "[WILDCARD]  ", bold=True, color=AMBER, size=9)
    _emit(p, str(q.get("text", "")), INK, 10.5)
    mk = q.get("marks")
    if mk not in (None, ""):
        _run(p, f"   [{mk} mark{'s' if str(mk) != '1' else ''}]", bold=True, color=GREY, size=9)
    _diagram_boxes(doc, q.get("diagram_prompts"))
    _embed_image(doc, q.get("image"))
    for opt in q.get("options") or []:
        op = _para(doc, space_after=1, indent=0.3)
        _emit(op, str(opt), INK, 10.5)


def _is_rich(sol):
    return bool(sol.get("concept") or sol.get("where") or sol.get("option_analysis"))


def render_answer_concept_solution(doc, q, show_solutions):
    # Answer key
    p = _para(doc, space_before=3, space_after=1, indent=0.2)
    _run(p, "Answer Key:  ", bold=True, color=GREEN, size=10)
    _emit(p, str(q.get("answer", "")), INK, 10, bold=True)
    sol = q.get("solution") or {}
    rich = _is_rich(sol)
    # legacy one-line concept (only when the solution isn't the rich form)
    if not rich and q.get("concept"):
        p = _para(doc, space_after=1, indent=0.2)
        _run(p, "Concept:  ", bold=True, color=TEAL, size=10)
        _emit(p, str(q["concept"]), INK, 10)
    if not show_solutions or not sol:
        return
    if rich:
        _render_rich_solution(doc, q, sol)
    else:
        _render_legacy_solution(doc, sol)
    _diagram_boxes(doc, sol.get("solution_diagram_prompts"))


def _subhead(doc, text):
    p = _para(doc, space_before=3, space_after=1, indent=0.2)
    _run(p, text, bold=True, color=TEAL, size=10)


def _render_rich_solution(doc, q, sol):
    p = _para(doc, space_before=3, space_after=1, indent=0.2)
    _run(p, "Detailed Solution", bold=True, color=NAVY, size=10.5)

    # 1. Core Concept & Formula
    _subhead(doc, "1. Core Concept & Formula")
    concept = sol.get("concept") or q.get("concept")
    if concept:
        cp = _para(doc, space_after=2, indent=0.35)
        _emit(cp, str(concept), INK, 10)
    for f in sol.get("formula") or []:
        fp = _para(doc, space_after=1, indent=0.55)
        _emit(fp, str(f), INK, 10.5, whole_math=True)
    where = sol.get("where") or []
    if where:
        wp = _para(doc, space_before=1, space_after=0, indent=0.35)
        _run(wp, "Where:", bold=True, color=NAVY, size=9.5)
        for g in where:
            lp = _para(doc, space_after=0, indent=0.65)
            _run(lp, "•  ", bold=True, color=TEAL, size=9.5)
            sym = str(g.get("sym", "")).strip()
            if sym:
                _emit(lp, sym, INK, 9.5, bold=True, whole_math=True)
                _run(lp, " = ", color=INK, size=9.5)
            _emit(lp, str(g.get("def", "")), INK, 9.5)  # def may carry symbols/units -> render math
            if g.get("unit"):
                _run(lp, f"  ({g['unit']})", color=INK, size=9.5)

    # 2. Calculation Steps  OR  Evaluating the Options
    opts = sol.get("option_analysis") or []
    calc = sol.get("calculation") or []
    if opts:
        _subhead(doc, "2. Evaluating the Options")
        for o in opts:
            lp = _para(doc, space_after=1, indent=0.45)
            correct = o.get("correct")
            _run(lp, f"({o.get('option', '')}) ", bold=True,
                 color=(GREEN if correct else RED), size=9.5)
            _emit(lp, str(o.get("text", "")), INK, 9.5)
    elif calc:
        _subhead(doc, "2. Calculation Steps")
        for line in calc:
            lp = _para(doc, space_after=0, indent=0.55)
            _emit(lp, str(line), INK, 9.5, whole_math=True)

    fa = sol.get("final_answer")
    if fa:
        sp = _para(doc, space_before=2, space_after=1, indent=0.35)
        _run(sp, "Final Answer:  ", bold=True, color=GREEN, size=10)
        _emit(sp, str(fa), INK, 10, bold=True)


def _render_legacy_solution(doc, sol):
    p = _para(doc, space_after=1, indent=0.2)
    _run(p, "Solution:", bold=True, color=NAVY, size=10)
    _solution_step(doc, "Step 1 — Given Data", sol.get("given"))
    _solution_step(doc, "Step 2 — Governing Formula", sol.get("formula"), whole_math=True)
    _solution_step(doc, "Step 3 — Line-by-Line Calculation", sol.get("calculation"), whole_math=True)
    fa = sol.get("final_answer")
    if fa:
        sp = _para(doc, space_after=1, indent=0.45)
        _run(sp, "Step 4 — Final Answer:  ", bold=True, color=GREEN, size=9.5)
        _emit(sp, str(fa), INK, 9.5, bold=True)


def _solution_step(doc, label, payload, whole_math=False):
    if payload in (None, "", []):
        return
    sp = _para(doc, space_after=0, indent=0.45)
    _run(sp, f"{label}:", bold=True, color=NAVY, size=9.5)
    items = payload if isinstance(payload, list) else [payload]
    for it in items:
        lp = _para(doc, space_after=0, indent=0.65)
        _emit(lp, str(it), INK, 9.5, whole_math=whole_math)


def render_question_combined(doc, q, show_solutions, tags):
    render_stem(doc, q)
    # Per-question classification now lives in the separate metadata .xlsx, not inline.
    render_answer_concept_solution(doc, q, show_solutions)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def render_question_only(doc, q, tags):
    render_stem(doc, q)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ----------------------------- answer-key table (split layout) ------------
def add_answer_key_table(doc, questions):
    p = _para(doc, space_before=4, space_after=3)
    _run(p, "Quick Answer Key", bold=True, color=NAVY, size=13)
    ncols = 4
    rows = (len(questions) + ncols - 1) // ncols
    table = doc.add_table(rows=rows, cols=ncols * 2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, q in enumerate(questions):
        r = idx % rows
        c = (idx // rows) * 2
        qcell = table.cell(r, c)
        acell = table.cell(r, c + 1)
        _shade(qcell, FILL_TINT_TEAL)
        qp = qcell.paragraphs[0]
        qp.paragraph_format.space_after = Pt(0)
        _run(qp, f"Q{q.get('number', '')}", bold=True, color=NAVY, size=9)
        ap = acell.paragraphs[0]
        ap.paragraph_format.space_after = Pt(0)
        _run(ap, str(q.get("answer", "")), color=INK, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ----------------------------- driver -------------------------------------
def build(spec, out_path, latex=False, mathsize=LX.DEFAULT_SIZE):
    P.normalize_paper(spec)        # coerce string formula/calc fields -> lists
    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    meta = spec.get("meta", {})
    layout = (meta.get("layout") or "combined").lower()
    show_solutions = meta.get("show_solutions", True)
    tags = meta.get("metadata_tags") or []
    sections = spec.get("sections", [])

    # Equation rendering: inline LaTeX (no pandoc) or native Word OMML.
    global _OMML_CACHE, LATEX_MODE, MATH_SIZE
    LATEX_MODE = latex
    if latex:
        MATH_SIZE = mathsize
        _OMML_CACHE = {}
        print(f"  math: inline LaTeX [{mathsize or 'full size'}] (no pandoc)")
    elif OM.available():
        _OMML_CACHE = OM.convert_batch(_collect_math(spec))
        print(f"  math: {sum(1 for v in _OMML_CACHE.values() if v)} equations -> native Word OMML")
    else:
        _OMML_CACHE = {}
        print("  math: pandoc unavailable -> Unicode fallback (install pypandoc_binary for OMML)")

    add_header(doc, meta)

    if layout == "split":
        # ---- Part 1: question paper (questions only) ----
        p = _para(doc, space_after=4)
        _run(p, "QUESTION PAPER", bold=True, color=TEAL, size=14)
        all_q = []
        for sec in sections:
            add_section_band(doc, sec.get("section", "Section"))
            for q in sec.get("questions", []):
                render_question_only(doc, q, tags)
                all_q.append(q)
        # ---- Part 2: answer key + solutions ----
        doc.add_page_break()
        p = _para(doc, space_after=4)
        _run(p, "ANSWER KEY & SOLUTIONS", bold=True, color=TEAL, size=14)
        add_answer_key_table(doc, all_q)
        _rule(doc)
        for sec in sections:
            add_section_band(doc, sec.get("section", "Section") + " — Solutions")
            for q in sec.get("questions", []):
                # re-show a compact stem so the solution stands alone
                sp = _para(doc, space_before=3, space_after=1)
                _run(sp, f"Q{q.get('number', '')}.  ", bold=True, color=TEAL, size=10.5)
                _emit(sp, str(q.get("text", "")), INK, 10)
                render_answer_concept_solution(doc, q, show_solutions)
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
    else:
        # ---- combined: everything under each question ----
        for sec in sections:
            add_section_band(doc, sec.get("section", "Section"))
            for q in sec.get("questions", []):
                render_question_combined(doc, q, show_solutions, tags)

    # footer
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, f"ExamDost — {meta.get('exam', 'Generated Paper')}", color=GREY, size=8)

    if latex:
        LX.fix_doc(doc)            # guarantee no $-span / number collisions
    LATEX_MODE = False             # reset module global for any later in-process build

    out_path = P.safe_output_path(out_path)
    doc.save(out_path)
    nq = sum(len(s.get("questions", [])) for s in sections)
    print(f"Wrote {out_path}  ({nq} questions, layout={layout}, solutions={'on' if show_solutions else 'off'})")


def main():
    args = [a for a in sys.argv[1:]]
    latex = False
    mathsize = LX.DEFAULT_SIZE
    if "--latex" in args:
        latex = True
        args.remove("--latex")
    if "--mathsize" in args:
        i = args.index("--mathsize")
        raw = args[i + 1]
        del args[i:i + 2]
        mathsize = "" if raw in ("", "none", "full") else ("\\" + raw.lstrip("\\"))
    if len(args) != 2:
        print(__doc__)
        sys.exit(1)
    with open(args[0], encoding="utf-8") as f:
        spec = json.load(f)
    build(spec, args[1], latex=latex, mathsize=mathsize)


if __name__ == "__main__":
    main()
