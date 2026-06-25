r"""
build_paper_boxed_docx.py — render a paper (paper.json) as the BOXED / field-table Word
format (one two-column table per question), for content-team / platform import.

Usage:
    python build_paper_boxed_docx.py <paper.json> <out.docx> [--latex] [--mathsize <s>]

    --latex            Emit equations as inline LaTeX ($\scriptstyle ..$) text instead of
                       native Word OMML — the upload-ready form, no pandoc, no OMML->LaTeX
                       conversion step. (Default: native Word equations / OMML.)
    --mathsize <s>     LaTeX size wrapper for --latex: scriptstyle (default) | scriptscriptstyle
                       | "" (full size). Accepts the bare word or a leading backslash.

Each question becomes a bordered table:
    Question | <n>
    Type     | MCQ / MSQ / NAT
    Body     | <question text + [GEMINI_FLASH_PROMPT: ...] image prompts>
    Option   | <opt1 value only — no A/B/C/D label>   (one row per option; omitted for NAT)
    Option   | <opt2> ...
    Correct  | <option number(s) for MCQ/MSQ, or numeric value/range for NAT>
    Explanation | <full rich solution: Core Concept & Formula -> steps/options -> Final>
    Subject  | <subject>
    Topic    | <Chapter Name (Type Name)>
    Correct Marks | <n>
    Incorrect Marks | <n, magnitude — no minus sign>
    Hint | -
    Video Solution | -
    PYQ |
    OTS ID |

Math renders as native Word equations (OMML); images are shown as labelled Gemini
prompts. Reuses the docx engine in build_paper_docx (_emit / _run / OMML cache).

Dependencies: python-docx, plus sibling build_paper_docx, paperlib, omml.
"""
from __future__ import annotations

import json
import os
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import paperlib as P
import omml as OM
import latexmode as LX
import build_paper_docx as BD

TYPE_FIELDS = ("Hint", "Video Solution", "PYQ", "OTS ID")


def boxed_type(q):
    t = (q.get("type") or "").strip().upper()
    if t == "NAT":
        return "NAT"
    if t == "MSQ":
        return "MSQ"
    return "MCQ"          # MCQ, Statement, Assertion-Reason all present as single-best MCQ


def correct_value(q):
    """Option number(s) for MCQ/MSQ; numeric value/range for NAT."""
    t = boxed_type(q)
    if t == "NAT":
        return str(q.get("answer", "")).strip()
    letters = P.answer_letters(q.get("answer", ""))
    nums = [str(ord(c.upper()) - 64) for c in letters]  # A->1, B->2, ...
    return ", ".join(nums)


def _fmt_marks(x):
    """Marks as a clean magnitude string (no minus sign, no trailing .0)."""
    try:
        v = abs(float(x))
    except (TypeError, ValueError):
        return str(x)
    return str(int(v)) if v == int(v) else f"{v:g}"


def marks_pair(q, meta):
    """Resolve (correct, incorrect) marks. Priority: per-question fields ->
    meta.marking flat values ('correct'/'incorrect') -> MCQ fraction default."""
    mrk = meta.get("marking") or {}
    cm = q.get("correct_marks")
    if cm is None:
        cm = mrk.get("correct", q.get("marks"))
    im = q.get("incorrect_marks")
    if im is None:
        if "incorrect" in mrk:
            im = mrk["incorrect"]
        else:
            try:
                mk = float(q.get("marks") or 0)
            except (TypeError, ValueError):
                mk = 0
            im = round(mk * float(mrk.get("mcq_neg_fraction", 1.0 / 3.0)), 2) if boxed_type(q) == "MCQ" else 0
    return _fmt_marks(cm), _fmt_marks(im)


# ----------------------------- cell helpers (reuse BD engine) --------------
def _cpar(cell, first=False, space_after=2, space_before=0, indent=None):
    if first and len(cell.paragraphs) == 1 and not cell.paragraphs[0].runs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.05
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def _image_prompt(cell, tag):
    p = _cpar(cell, space_after=2)
    BD._run(p, "[ Image — Gemini prompt ]  ", bold=True, color=BD.AMBER, size=9)
    BD._run(p, P.gemini_prompt(tag), italic=True, color=BD.INK, size=9)


def _body_cell(cell, q):
    p = _cpar(cell, first=True, space_after=2)
    BD._emit(p, str(q.get("text", "")), BD.INK, 10.5)
    img = q.get("image")
    if img and os.path.exists(img):
        try:
            _cpar(cell, space_after=2).add_run().add_picture(img, width=Inches(3.6))
        except Exception:
            pass
    for tag in q.get("diagram_prompts") or []:
        _image_prompt(cell, tag)


def _rich_solution_cell(cell, q, sol):
    # 1. Core Concept & Formula
    h = _cpar(cell, first=True, space_after=1)
    BD._run(h, "1. Core Concept & Formula", bold=True, color=BD.TEAL, size=10)
    concept = sol.get("concept") or q.get("concept")
    if concept:
        BD._emit(_cpar(cell, space_after=2), str(concept), BD.INK, 10)
    for f in sol.get("formula") or []:
        BD._emit(_cpar(cell, space_after=1, indent=0.15), str(f), BD.INK, 10.5, whole_math=True)
    where = sol.get("where") or []
    if where:
        BD._run(_cpar(cell, space_after=0), "Where:", bold=True, color=BD.NAVY, size=9.5)
        for g in where:
            lp = _cpar(cell, space_after=0, indent=0.3)
            BD._run(lp, "•  ", bold=True, color=BD.TEAL, size=9.5)
            sym = str(g.get("sym", "")).strip()
            if sym:
                BD._emit(lp, sym, BD.INK, 9.5, bold=True, whole_math=True)
                BD._run(lp, " = ", color=BD.INK, size=9.5)
            BD._emit(lp, str(g.get("def", "")), BD.INK, 9.5)  # def may carry symbols -> render math
            if g.get("unit"):
                BD._run(lp, f"  ({g['unit']})", color=BD.INK, size=9.5)
    # 2. Calculation Steps / Evaluating the Options
    opts = sol.get("option_analysis") or []
    calc = sol.get("calculation") or []
    if opts:
        BD._run(_cpar(cell, space_before=2, space_after=1), "2. Evaluating the Options",
                bold=True, color=BD.TEAL, size=10)
        for o in opts:
            lp = _cpar(cell, space_after=1, indent=0.15)
            BD._run(lp, f"({o.get('option', '')}) ", bold=True,
                    color=(BD.GREEN if o.get("correct") else BD.RED), size=9.5)
            BD._emit(lp, str(o.get("text", "")), BD.INK, 9.5)
    elif calc:
        BD._run(_cpar(cell, space_after=1), "2. Calculation Steps", bold=True, color=BD.TEAL, size=10)
        for line in calc:
            BD._emit(_cpar(cell, space_after=0, indent=0.15), str(line), BD.INK, 9.5, whole_math=True)
    fa = sol.get("final_answer")
    if fa:
        fp = _cpar(cell, space_before=2, space_after=1)
        BD._run(fp, "Final Answer:  ", bold=True, color=BD.GREEN, size=10)
        BD._emit(fp, str(fa), BD.INK, 10, bold=True)
    for tag in sol.get("solution_diagram_prompts") or []:
        _image_prompt(cell, tag)


def _label(cell, text):
    BD._shade(cell, "EEF1F6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    BD._run(p, text, bold=True, color=BD.NAVY, size=9.5)


def _plain(cell, text):
    p = _cpar(cell, first=True, space_after=0)
    BD._emit(p, str(text), BD.INK, 10)


def _size_math_runs(doc, fallback="21"):
    """Stamp an explicit font size on every native-OMML math run.

    Word emits OMML math runs (<m:r>) with NO <w:sz>, so equations render at Word's
    default math size — visibly SMALLER than the 10.5pt body text around them. For
    each math run we set <w:sz>/<w:szCs> to the largest text-run size in its OWN
    paragraph (fallback 10.5pt = sz 21), so every equation matches the prose beside
    it. No-op in inline-LaTeX mode (no <m:oMath> elements). Returns runs sized."""
    W_SZ, W_SZCS, W_RPR, M_RPR = qn("w:sz"), qn("w:szCs"), qn("w:rPr"), qn("m:rPr")
    M_OMATH, M_R, W_R, W_VAL = qn("m:oMath"), qn("m:r"), qn("w:r"), qn("w:val")

    def set_run_size(mr, val):
        # ensure <w:rPr> exists in the math run, after <m:rPr> if present, before <m:t>
        wrpr = mr.find(W_RPR)
        if wrpr is None:
            wrpr = mr.makeelement(W_RPR, {})
            mrpr = mr.find(M_RPR)
            if mrpr is not None:
                mrpr.addnext(wrpr)
            else:
                mr.insert(0, wrpr)
        for tag in (W_SZ, W_SZCS):
            el = wrpr.find(tag)
            if el is None:
                el = wrpr.makeelement(tag, {}); wrpr.append(el)
            el.set(W_VAL, val)

    def target_size(p):
        szs = []
        for r in p.findall(W_R):
            rpr = r.find(W_RPR)
            if rpr is not None:
                sz = rpr.find(W_SZ)
                if sz is not None and sz.get(W_VAL):
                    szs.append(int(sz.get(W_VAL)))
        return str(max(szs)) if szs else fallback

    n = 0
    for p in doc.element.body.iter(qn("w:p")):
        omaths = p.findall(".//" + M_OMATH)
        if not omaths:
            continue
        tgt = target_size(p)
        for om in omaths:
            for mr in om.iter(M_R):
                set_run_size(mr, tgt); n += 1
    return n


def build(spec, out_path, latex=False, mathsize=LX.DEFAULT_SIZE):
    P.normalize_paper(spec)        # coerce string formula/calc fields -> lists
    out_path = P.safe_output_path(out_path)
    meta = spec.get("meta", {})
    sections = spec.get("sections", [])

    # equation rendering: inline LaTeX (no pandoc) or native Word OMML
    BD.LATEX_MODE = latex
    if latex:
        BD.MATH_SIZE = mathsize
        BD._OMML_CACHE = {}
    else:
        # pre-convert all equations to OMML once
        BD._OMML_CACHE = OM.convert_batch(BD._collect_math(spec)) if OM.available() else {}

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    h = doc.add_paragraph()
    BD._run(h, meta.get("paper_title") or meta.get("exam", "Question Bank"),
            bold=True, color=BD.NAVY, size=15)

    nq = 0
    for sec in sections:
        for q in sec.get("questions", []):
            nq += 1
            rows = [("Question", "plain", str(q.get("number", nq))),
                    ("Type", "plain", boxed_type(q)),
                    ("Body", "body", q)]
            if boxed_type(q) != "NAT":
                for opt in q.get("options") or []:
                    rows.append(("Option", "plain", P.option_value(opt)))
            rows.append(("Correct", "plain", correct_value(q)))
            rows.append(("Explanation", "expl", q))
            cm, im = marks_pair(q, meta)
            rows += [
                ("Subject", "plain", q.get("subject", "")),
                ("Topic", "plain", f"{q.get('chapter','')} ({q.get('topic','')})"),
                ("Correct Marks", "plain", cm),
                ("Incorrect Marks", "plain", im),
                ("Hint", "plain", "-"),
                ("Video Solution", "plain", "-"),
                ("PYQ", "plain", ""),
                ("OTS ID", "plain", ""),
            ]
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for label, kind, val in rows:
                r = table.add_row()
                c0, c1 = r.cells[0], r.cells[1]
                c0.width = Inches(1.9)
                c1.width = Inches(5.1)
                _label(c0, label)
                if kind == "body":
                    _body_cell(c1, val)
                elif kind == "expl":
                    sol = val.get("solution") or {}
                    if sol:
                        _rich_solution_cell(c1, val, sol)
                    else:
                        _plain(c1, "")
                else:
                    _plain(c1, val)
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    if latex:
        fixes = LX.fix_doc(doc)   # guarantee no $-span / number collisions
    else:
        _size_math_runs(doc)      # equations render at body-text size, not Word's smaller default
    BD.LATEX_MODE = False         # reset module global for any later build in-process

    doc.save(out_path)
    mode = f"inline LaTeX [{mathsize or 'full size'}]" if latex else "native Word OMML"
    print(f"Wrote {out_path}  ({nq} questions, boxed field-table, {mode})")


def main():
    args = [a for a in sys.argv[1:]]
    latex = False
    mathsize = LX.DEFAULT_SIZE
    if "--latex" in args:
        latex = True
        args.remove("--latex")
    if "--mathsize" in args:
        i = args.index("--mathsize")
        raw = args[i + 1]
        del args[i:i + 2]
        mathsize = "" if raw in ("", "none", "full") else ("\\" + raw.lstrip("\\"))
    if len(args) != 2:
        print(__doc__)
        sys.exit(1)
    with open(args[0], encoding="utf-8") as f:
        spec = json.load(f)
    build(spec, args[1], latex=latex, mathsize=mathsize)


if __name__ == "__main__":
    main()
